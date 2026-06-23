from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "实验报告.docx"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_cell_margins(table, top=80, bottom=80, start=120, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    values = {"top": top, "bottom": bottom, "start": start, "end": end}
    for side, value in values.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, after=6, before=0, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(doc, text, level):
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color="2E74B5")
        style_paragraph(paragraph, before=16, after=8, line=1.10)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color="2E74B5")
        style_paragraph(paragraph, before=12, after=6, line=1.10)
    else:
        set_run_font(run, size=12, bold=True, color="1F4D78")
        style_paragraph(paragraph, before=8, after=4, line=1.10)
    return paragraph


def add_body(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=8, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=8, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_cell_margins(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("分布式黑板多车探索实验报告")
    set_run_font(run, size=9, color="666666")


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("分布式黑板风格多车协作探索系统实验报告")
    set_run_font(run, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("基于 Java、Redis、RabbitMQ 与现有 Swing View 的实现方案")
    set_run_font(run, size=12, color="1F4D78")

    add_heading(doc, "一、实验目标", 1)
    add_body(doc, "本实验要求设计并实现一个分布式黑板风格的多车协作地图探索系统。系统将地图抽象为二维网格，小车作为知识源在控制器节拍驱动下移动，并通过 Redis 黑板共享地图、障碍物、车辆状态和路径数据。")
    add_body(doc, "本次实现保留已有 Swing View 作为界面组件，在其基础上新增 InspectionBackend 后端工程，补齐控制器、导航器、小车知识源、记录回放和初始化工具。")

    add_heading(doc, "二、需求对照", 1)
    add_table(
        doc,
        ["实验要求", "实现方式", "验收效果"],
        [
            ["共享数据", "Redis DB9 保存地图、小车、路径、用户和回放帧", "各模块不直接持有彼此状态"],
            ["组件独立", "Controller、Navigator、Car、Recorder 均有独立 main", "可一键启动，也可分进程部署"],
            ["消息调度", "RabbitMQ fanout/direct 交换机传递开始、节拍和导航任务", "控制流通过消息总线完成"],
            ["地图探索", "小车移动后点亮当前位置周围 3x3 区域", "界面能看到探索区域持续增加"],
            ["路径规划", "支持 A*、双向 A*、Dijkstra 三种算法", "配置员可在 View 中切换算法"],
            ["可视化与回放", "沿用 View 实时地图和 Analyst 回放界面", "可查看小车位置、障碍物、探索过程和历史记录"],
        ],
        [1900, 3950, 3510],
    )

    add_heading(doc, "三、总体架构", 1)
    add_body(doc, "系统采用黑板风格。Redis 是黑板，保存所有共享数据；RabbitMQ 是消息总线，承担控制器与知识源之间的消息传递；Controller 是唯一调度器，Navigator、Car、Recorder 是相互独立的知识源；View 是显示和配置终端。")
    add_bullet(doc, "Controller：监听 controller.start，收到开始命令后按固定节拍循环，分配导航任务并广播小车移动命令。")
    add_bullet(doc, "Navigator：监听 navigator.no1 到 navigator.no4，读取地图和算法配置，规划路径并写入小车任务队列。")
    add_bullet(doc, "Car：自动发现 Cars:1 到 Cars:5，收到节拍后读取路径队列，移动一步并点亮地图。")
    add_bullet(doc, "Recorder：监听 save.start，保存探索帧；分析员回放时按 Save 字段恢复历史帧。")

    add_heading(doc, "四、黑板数据设计", 1)
    add_table(
        doc,
        ["Redis Key", "类型", "说明"],
        [
            ["map_width / map_height", "String", "地图宽高，由 View 配置界面写入"],
            ["MapView", "Bitmap", "地图探索状态，1 表示已探索"],
            ["blockview", "Bitmap", "障碍物状态，1 表示存在障碍"],
            ["Cars:1..5", "Hash", "小车 x、y、endx、endy、state、direction"],
            ["1_task_queue..5_task_queue", "List", "小车路径队列，元素为 x,y"],
            ["Algorithm", "String", "0=A*，1=双向 A*，2=Dijkstra"],
            ["Save", "Hash", "回放记录数量、当前记录、速度和进度"],
            ["Record:<file>:<frame>", "Hash", "记录器保存的回放帧快照"],
        ],
        [2500, 1500, 5360],
    )

    add_heading(doc, "五、核心运行流程", 1)
    for item in [
        "配置员在 View 中设置地图、小车、障碍物和路径算法。",
        "View 点击开始后向 controller 和 save 交换机发送启动消息。",
        "Controller 扫描 Redis 中空闲小车，将其当前位置封装为导航任务。",
        "Navigator 选择未探索目标点，执行路径规划，将路径写入对应任务队列。",
        "Controller 广播节拍，小车收到节拍后移动一步并更新 Redis。",
        "View 周期读取 Redis，实时绘制小车、障碍物和已探索区域。",
        "Recorder 持续保存快照；分析员进入回放后，Recorder 按历史帧恢复 Redis 状态。",
    ]:
        add_number(doc, item)

    add_heading(doc, "六、关键算法", 1)
    add_body(doc, "导航器通过 PathFinder 接口封装算法，当前实现 A*、双向 A* 和 Dijkstra。算法均基于网格邻接关系，避开 blockview 中的障碍点，并把其他小车当前位置视为动态障碍；为了鼓励探索，已探索点的移动代价高于未探索点。")
    add_body(doc, "目标点选择采用广度优先扫描：优先选择距离当前车辆不太近也不太远的未探索点，若不存在理想点，则从剩余未探索点中选择 fallback 目标。")

    add_heading(doc, "七、测试与验收", 1)
    add_table(
        doc,
        ["测试项", "检查方法", "期望结果"],
        [
            ["路径算法", "构造有障碍地图运行单元测试", "三种算法均返回不穿越障碍的合法路径"],
            ["3x3 点亮", "测试中心点和角落点", "中心点点亮 9 格，角落点点亮 4 格"],
            ["消息拓扑", "运行 tools init", "RabbitMQ 队列和交换机声明成功"],
            ["实时探索", "启动 launcher 后在 View 点击开始", "小车移动，MapView bitcount 持续增长"],
            ["回放", "完成一次探索后进入 Analyst", "记录列表出现，进度和倍速可控制"],
        ],
        [2000, 3650, 3710],
    )

    add_heading(doc, "八、运行方式", 1)
    add_body(doc, "首先启动 Redis 和 RabbitMQ，然后在 InspectionBackend 目录执行 Maven 构建和初始化命令。初始化会写入 admin、config、analyst 三个默认用户，并声明所有消息队列。")
    add_body(doc, "后端可通过 launcher 一键启动，也可以分别启动 controller、navigator、car、recorder 四个模块以展示分布式部署能力。最后启动现有 View 项目中的 Login.Main，用 config/config123 登录并开始探索。")

    add_heading(doc, "九、结论", 1)
    add_body(doc, "该实现满足实验文档中对分布式黑板风格的核心要求：组件通过 Redis 共享数据，通过 RabbitMQ 交换控制消息，控制器统一节拍调度，知识源可独立部署，且系统具备可视化运行状态和回放能力。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
